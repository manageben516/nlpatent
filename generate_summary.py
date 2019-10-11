import docx, codecs
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from docx.shared import Pt
from gen_const import DIR, CPC_PATH, TITLE_PATH, ABSTRACT_PATH
from gen_const import generate_daystamp, get_titles, get_abstracts, get_CPC_descriptions

##### HELPERS #####

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def add_styles(styles, txtFont, txtSize):
    # this changes the normal text font & size
    style = styles['Normal']
    font = style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    para_format = style.paragraph_format
    para_format.space_after = Pt(0)
    # OTHER OPTIONS: font.bold, font.italic, font.underline = True, True, True

    # this changes the style of 'Quote'
    style = styles['Quote']
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # or LEFT/RIGHT/JUSTIFY

    # this defines a new style for the Title (bug: can't change style when doing .add_heading)
    new_title_style = styles.add_style('NewTitle', WD_STYLE_TYPE.PARAGRAPH)
    new_title_style.base_style = styles['Title']
    font = new_title_style.font
    font.name = txtFont
    font.size = Pt(26)

    # this defines a new style for Heading 1 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading1', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+3)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 2 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading2', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+2)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 3 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading3', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 3']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+1)
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 4 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading4', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 4']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    font.italic = True
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for the Input Text
    new_style = styles.add_style('InputText', WD_STYLE_TYPE.PARAGRAPH)
    new_style.base_style = styles['Normal']
    font = new_style.font
    font.size = Pt(9)

def add_saved_prior_art(pa_saved, titles, abstracts, document, multiple_subclasses=False):
    p = document.add_paragraph('Saved Prior Art', style='NewHeading3')
    run = p.add_run()
    run.add_break()
    # display the additional prior art
    count = 1
    for p_n, cpc in pa_saved:
        p = document.add_paragraph()
        p.add_run('{}.  '.format(count))
        # if it's a publication number
        if len(p_n) > 8:
            add_hyperlink(p, 'https://patents.google.com/patent/US{}'.format(p_n), 'US {}'.format("{}/{}".format(p_n[:4], p_n[4:])), 'FF8822', True)
            if multiple_subclasses: p.add_run(u' | {}'.format(cpc)).italic = True
            run = p.add_run()
            run.add_break()
            p.add_run(u'{}'.format(titles[p_n].title())).bold = True # change from all UPPER to title case
            run = p.add_run()
            run.add_break()
            p = document.add_paragraph(u'{}'.format(abstracts[p_n]))
        # if it's a grant number
        else:
            add_hyperlink(p, 'https://patents.google.com/patent/US{}'.format(p_n), 'US {}'.format("{:,}".format(int(p_n))), 'FF8822', True)
            if multiple_subclasses: p.add_run(u' | {}'.format(cpc)).italic = True
            run = p.add_run()
            run.add_break()
            p.add_run(u'{}'.format(titles[p_n])).bold = True
            run = p.add_run()
            run.add_break()
            p = document.add_paragraph(u'{}'.format(abstracts[p_n]))
        count +=1
        p = document.add_paragraph()

def add_disclaimer(document):
    p = document.add_paragraph('DISCLAIMER', style='NewHeading1')
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break()
    p = document.add_paragraph('This summary has been generated automatically.')
    run = p.add_run()
    run.add_break()
    p = document.add_paragraph('As such, Legalicity provides no warranty and makes no claims regarding the accuracy, completeness, reliability, and overall quality of the content.')
    run = p.add_run()
    run.add_break()
    p = document.add_paragraph('This summary does not constitute legal or professional advice, and Legalicity shall not be liable for any loss or damage arising out of, or in connection with, any reliance on the contents of this summary.')
    run = p.add_run()
    run.add_break()

##### SECONDARY #####

def run(pa_saved, user_name, c_id, email, project_name):
    # CREATE A DOCUMENT & SET THE STYLES
    document = Document()
    styles = document.styles
    add_styles(styles, 'Arial', 11)

    # ADD TITLE TO DOCUMENT
    document.add_paragraph('Summary of NLPatent Results', style='NewTitle')

    # ADD LOGO TO DOCUMENT
    p = document.add_paragraph()
    run = p.add_run()
    run.add_picture('{}/clients/{}/logo.jpg'.format(DIR, c_id))

    # GENERATE TIMESTAMP
    p_l = generate_daystamp()

    # ADD TIMESTAMP & OWNER TO DOCUMENT
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('{}'.format(p_l)).bold = True
    run = p.add_run()
    run.add_break()
    p.add_run('Prepared for user: ')
    p.add_run(u'{}'.format(user_name)).italic = True

    # ADD PROJECT TITLE TO DOCUMENT
    p = document.add_paragraph('Project Title', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(project_name)).bold = True

    # get a sorted list of CPC subclasses
    cpc_subclasses = sorted(set(list(zip(*pa_saved)[1])))
    # get a sorted list of CPC sections
    cpc_sections = sorted(set([x[0] for x in cpc_subclasses]))
    # get the CPC descriptions
    cpc_descr = get_CPC_descriptions('{}/{}/cpc-descriptions.txt'.format(DIR, CPC_PATH))

    # ADD TECHNOLOGY AREA TO DOCUMENT
    if len(cpc_subclasses) > 1:
        p = document.add_paragraph('CPC Categories of the Results', style='NewHeading3')
    else:
        p = document.add_paragraph('CPC Category of the Results', style='NewHeading3')
    # go through every CPC section
    for s in cpc_sections:
        p = document.add_paragraph()
        # Add one section, e.g. 'A'
        p.add_run('{}: '.format(s)).bold = True
        p.add_run('{}'.format(cpc_descr[s]))
        run = p.add_run()
        run.add_break()
        # Add classes & subclasses
        # get only the subclasses for this section
        s_sc = [x for x in cpc_subclasses if x[0] == s]
        # get only the classes for this section
        s_c = sorted(set([x[:-1] for x in s_sc]))
        # go through class by class
        for c in s_c:
            if len(s_c) > 1:
                run = p.add_run()
                run.add_break()
            p.add_run('{}: '.format(c)).bold = True
            p.add_run('{}'.format(cpc_descr[c]))
            run = p.add_run()
            run.add_break()
            # go through each class' subclasses
            for sc in s_sc:
                if sc[:-1] == c:
                    p.add_run('{}: '.format(sc)).bold = True
                    p.add_run('{}'.format(cpc_descr[sc]))
                    run = p.add_run()
                    run.add_break()

    # ADD INPUT TEXT TO DOCUMENT
    # read in the text of the patent document
    with codecs.open('{}/clients/{}/{}/{}/input_text.txt'.format(DIR, c_id, email, project_name), 'r', 'utf-8') as f_in:
        l = f_in.read()
    p = document.add_paragraph('Input Text', style='NewHeading3')
    p = document.add_paragraph('The following is a copy of the input text:')
    run = p.add_run()
    run.add_break()
    p = document.add_paragraph(l, style='InputText') # input text is smaller - 9pt

    document.add_page_break() # start on new page AFTER the input text

    # ADD SAVED PRIOR ART TO DOCUMENT
    # first, get the titles of ONLY the prior art
    titles, temp_titles = {}, {}
    # go through each chosen subclass
    for cpc in cpc_subclasses:
        temp_titles = get_titles('{}/{}/{}.txt'.format(DIR, TITLE_PATH, cpc))
        titles.update({a: temp_titles[a] for a, b in pa_saved if b == cpc})
    temp_titles = {} # don't need to store anymore
    # then, get the abstracts of ONLY the prior art
    abstracts, temp_abstracts = {}, {}
    # go through each chosen subclass
    for cpc in cpc_subclasses:
        temp_abstracts = get_abstracts('{}/{}/{}.txt'.format(DIR, ABSTRACT_PATH, cpc))
        abstracts.update({a: temp_abstracts[a] for a, b in pa_saved if b == cpc})
    temp_abstracts = {} # don't need to store anymore
    # finally, add the prior art
    run = p.add_run()
    run.add_break()
    add_saved_prior_art(pa_saved, titles, abstracts, document, multiple_subclasses=len(cpc_subclasses)>1)

    document.add_page_break() # start on new page AFTER the saved prior art

    # ADD DISCLAIMER
    add_disclaimer(document)

    # SAVE DOCUMENT
    f_name = '[{}] {}.docx'.format(p_l, project_name)
    document.save('{}/clients/{}/{}/{}/{}'.format(DIR, c_id, email, project_name, f_name))

    return f_name
