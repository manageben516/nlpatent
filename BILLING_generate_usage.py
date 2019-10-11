# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import Pt
import os, codecs

##### CONSTANTS ######

DIR = '/legalicity/nlpatent'
BILLING_PATH = '/legalicity/nlpatent/billing'

##### HELPERS #####

def add_styles(styles, txtFont, txtSize):
    # this changes the normal text font & size
    style = styles['Normal']
    font = style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    para_format = style.paragraph_format
    para_format.space_after = Pt(0)
    # OTHER OPTIONS: font.bold, font.italic, font.underline = True, True, True

    # this changes the style of 'Quote'
    style = styles['Quote']
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # or LEFT/RIGHT/JUSTIFY

    # this defines a new style for the Title (bug: can't change style when doing .add_heading)
    new_title_style = styles.add_style('NewTitle', WD_STYLE_TYPE.PARAGRAPH)
    new_title_style.base_style = styles['Title']
    font = new_title_style.font
    font.name = txtFont
    font.size = Pt(26)

    # this defines a new style for Heading 1 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading1', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+3)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 2 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading2', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+2)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 3 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading3', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 3']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+1)
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 4 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading4', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 4']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    font.italic = True
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for the Input Text
    new_style = styles.add_style('InputText', WD_STYLE_TYPE.PARAGRAPH)
    new_style.base_style = styles['Normal']
    font = new_style.font
    font.size = Pt(9)

def get_usage(c_id, email):
    p_n, p_u, t_u = [], {}, (0, 0)
    if os.path.exists('{}/clients/{}/{}/usage.txt'.format(DIR, c_id, email)):
        # get all the lines in the usage file
        lines = [line.strip() for line in codecs.open('{}/clients/{}/{}/usage.txt'.format(DIR, c_id, email), 'r', 'utf-8')]
        if lines:
            for l in lines:
                # extract the date, type (search or update), and project name
                date, t, p = l.split()[0], l.split()[1], ' '.join(l.split()[2:])
                # get the increment: add either one search or one update
                i = (1, 0) if t == 'S' else (0, 1)
                t_u = tuple(map(sum, zip(t_u, i)))
                # add a new project name to the list
                if p not in p_n: p_n.append(p)
                # add a new entry to the dict
                if p in p_u:
                    # check if this date already exists
                    if date in p_u[p]:
                        # get the existing number of searches & updates for this date
                        e_i = p_u[p][date]
                        # update the dict
                        p_u[p][date] = tuple(map(sum, zip(e_i, i))) # add the two tuples together
                    else:
                        p_u[p][date] = i
                else:
                    # add an entirely new entry
                    p_u[p] = {date: i}
    return p_n, p_u, t_u

def add_usage(document, d, c_id, s_cost, u_cost):
    p = document.add_paragraph('Users', style='NewHeading3')
    p = document.add_paragraph()
    # add each user's usage
    for x in sorted(d):
        # get the usage from usage.txt
        p_n, p_u, t_u = get_usage(c_id, x)
        # make sure the TOTAL number of searches & updates aligns
        if t_u[0] != d[x][0]:
            print('{} --> {} : {} searches in usage.txt but {} searches in MongoDB!'.format(c_id, x, t_u[0], d[x][0]))
            continue
        if t_u[1] != d[x][1]:
            print('{} --> {} : {} updates in usage.txt but {} updates in MongoDB!'.format(c_id, x, t_u[1], d[x][1]))
            continue
        # add the user's email address
        p.add_run(u'{}'.format(x)).italic = True
        p = document.add_paragraph()
        # add the by-project breakdowns
        for project_name in p_n:
            # get the number of searches & updates for this project
            p_t = map(sum, zip(*p_u[project_name].values()))
            # calculate project total
            project_total = p_t[0]*s_cost + p_t[1]*u_cost
            # add project name & total
            p.add_run(u'{}'.format(project_name)).underline = True
            p = document.add_paragraph()
            p.add_run(u'{}'.format('${:,.2f}'.format(project_total))).bold = True
            p = document.add_paragraph()
            # add the dates
            for date in sorted(p_u[project_name]):
                p.add_run(u'[{}] Searches: {}, Updates: {}'.format(date, p_u[project_name][date][0], p_u[project_name][date][1]))
                p = document.add_paragraph()
            p = document.add_paragraph()
        p = document.add_paragraph()

##### SECONDARY #####

def run(c_id, c, d, m, s_cost, u_cost):
    # CREATE A DOCUMENT & SET THE STYLES
    document = Document()
    styles = document.styles
    add_styles(styles, 'Arial', 11)

    # ADD TITLE
    document.add_paragraph('NLPatent Usage Summary', style='NewTitle')

    # ADD CLIENT ORGANIZATION
    p = document.add_paragraph('Client', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(c))

    # ADD USAGE MONTH
    p = document.add_paragraph('Month', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(m))

    # ADD CURRENT PRICING
    p = document.add_paragraph('Pricing', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'${} per search'.format(s_cost))
    p = document.add_paragraph()
    p.add_run(u'${} per update'.format(u_cost))

    # ADD PER USER USAGE SUMMARY
    add_usage(document, d, c_id, s_cost, u_cost)

    # SAVE DOCUMENT AS WORD FILE
    f_name = '{}_usage'.format(c_id)
    doc_name_word = '{}/{}/{}.docx'.format(BILLING_PATH, m, f_name)
    document.save(doc_name_word)

